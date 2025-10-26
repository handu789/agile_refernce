import csv
import re
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.shared import RGBColor
from docx.oxml import OxmlElement, ns
from datetime import datetime
client = OpenAI(api_key="<API_KEY>", base_url="<BASE_URL>")
# 读取模板和示例（只需一次）
with open("tech_example.txt", "r", encoding="utf-8") as example:
    example_content = example.read()

with open("tech_requirement.txt", "r", encoding="utf-8") as requirement:
    tech_requirement = requirement.read()


# 初始化文档设置
def init_document():
    doc = Document()
    section = doc.sections[0]

    # A4纸张设置
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # 页边距
    margins = section
    margins.left_margin = Cm(3.18)
    margins.right_margin = Cm(3.18)
    margins.top_margin = Cm(2.54)
    margins.bottom_margin = Cm(2.54)

    # 行网格设置
    sectPr = section._sectPr
    docGrid = OxmlElement('w:docGrid')
    docGrid.set(qn('w:type'), 'lines')
    docGrid.set(qn('w:linePitch'), '312')
    sectPr.append(docGrid)

    return doc


def add_hyperlink(paragraph, url, text, color="0000FF"):
    """ 添加仿宋字体的超链接 """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')

    # 创建格式属性
    rPr = OxmlElement('w:rPr')

    # 设置字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '仿宋')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:hAnsi'), '仿宋')
    rPr.append(rFonts)

    # 设置字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), "28")  # 14磅对应28 half-points
    rPr.append(sz)

    # 设置颜色和下划线
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


# 新增时间格式化函数
def format_time(original_time):
    # 去除可能的前导符号
    cleaned_time = original_time.lstrip('·').strip()
    # 使用正则匹配年月日
    match = re.search(r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})', cleaned_time)
    if match:
        year = int(match.group(1))
        month = int(match.group(2))
        day = int(match.group(3))
        return f"{year}年{month}月{day}日"
    # 若匹配失败返回原始值（可根据需要调整）
    return cleaned_time


# 处理单篇文章
def add_article(doc, answer):
    """ 根据结构化answer添加文章 """
    # 解包数据
    title = answer['metadata']['title']
    content = answer['content']
    time = answer['metadata']['time']
    href = answer['metadata']['href']

    # ===== 标题部分 =====
    title_para = doc.add_paragraph()
    # Wingdings符号
    symbol_run = title_para.add_run(chr(184))
    symbol_run.font.name = "Wingdings 2"
    symbol_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Wingdings 2')
    symbol_run.font.color.rgb = RGBColor(0, 0, 0)  # 新增颜色设置
    # 标题文本
    title_run = title_para.add_run(f" {title}")
    title_run.font.name = "仿宋"
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色
    # 段落格式设置
    title_para.style = 'Heading 2'  # 通过样式名称设置大纲级别
    title_para.paragraph_format.line_spacing = 1.16  # 新增行距设置
    title_para.paragraph_format.space_after = Pt(8)

    # ===== 正文处理 =====
    paragraphs = [p.strip() for p in content.split('\n') if p.strip()][:2]
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.16  # 明确设置行距
        p.paragraph_format.first_line_indent = Pt(28)  # 首行缩进
        run = p.add_run(para)
        run.font.name = "仿宋"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(8)

    # ===== 来源处理 =====
    # 时间段落（添加"时间："前缀）
    time_para = doc.add_paragraph(f"时间：{time}")  # 修改处
    time_para.paragraph_format.line_spacing = 1.16
    time_para.runs[0].font.name = "仿宋"
    time_para.runs[0].font.size = Pt(14)  # 新增字号设置
    time_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    time_para.paragraph_format.space_after = Pt(8)


    # 链接段落（带超链接）
    href_para = doc.add_paragraph()

    # 前缀部分
    prefix_run = href_para.add_run("链接：")
    prefix_run.font.name = "仿宋"
    prefix_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 显式设置中文字体
    prefix_run.font.size = Pt(14)

    # 超链接部分
    add_hyperlink(
        paragraph=href_para,
        url=href,
        text=href,
        color="0000FF"
    )

    # 设置段落格式
    href_para.paragraph_format.line_spacing = 1.16
    href_para.paragraph_format.space_after = Pt(8)

    # 强制刷新超链接字体设置
    for run in href_para.runs:
        if run.text.startswith("http"):
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(14)

    # 实时保存
    doc.save("xinzhiyuan.docx")


# 主处理流程
doc = init_document()

# 处理CSV文件
with open("xinzhiyuan.csv", "r", encoding="utf-8") as csvfile:
    csv_reader = csv.reader(csvfile)

    while True:
        try:
            # 读取表头行+数据行
            headers = next(csv_reader)
            data = next(csv_reader)
        except StopIteration:
            break  # 文件结束

        # 构建数据字典（跳过空行检查）
        if len(headers) == 4 and len(data) == 4:
            row = dict(zip(headers, data))

            # 格式化时间（新增处理步骤）
            formatted_time = format_time(row['time'])

            # 构建查询内容
            art_content = f"这是待处理的文章，请按模板处理：\n{row['passage']}\n。"

            # API调用
            response = client.chat.completions.create(
                model="YOUR_MODEL",
                messages=[
                    {"role": "system", "content": "你是一个专业的文本处理助理，能够完全理解并执行我提出的文本处理要求。"},
                    {"role": "user",
                     "content": f"我是一个科技类刊物的主编，我现在有一篇待处理的文章、三篇处理好的文章以及一个用于处理文章的要求模板，作为我的助理，你的任务是按照要求模板和示例来处理文章，并一定要保证处理结果的准确性。这是第三篇示例文章：\n{example_content}\n。"},
                    {"role": "assistant",
                     "content": "我已经学习了这三篇文章和对应的处理结果，我会学习每篇文章的处理方法，接下来请提供要求模板。"},
                    {"role": "user", "content": tech_requirement},  # 省略需求描述，可以从你的代码中复制需求部分
                    {"role": "assistant",
                     "content": "明白了，这是你给的处理要求模板，我将会严格遵循这个要求模板处理文章。接下来，请提供你需要我处理的文章，我会根据模板和示例的处理方式来进行处理，并严格保证处理结果的准确无误。"},
                    {"role": "user", "content": art_content}
                    # {"role": "user", "content": "在这之前，请复述一下三个示例的内容。"}
                ],
                temperature=1.1,
                top_p=0.9
            )
            # 构建结构化answer对象
            answer = {
                "metadata": {
                    "title": row['title'],
                    "time": formatted_time,
                    "href": row['href']
                },
                "content": response.choices[0].message.content,
                "formatted": f"""\
            ▌标题：{row['title']}
            ────────────────────────────
            ▌生成内容：
            {response.choices[0].message.content}
            ────────────────────────────
            ▌时间：{formatted_time}
            ▌链接：{row['href']}
            """
            }
            # 使用answer调用函数
            add_article(doc=doc,
                        answer=answer)  # 修改函数参数为接收answer
            # 使用格式化输出
            print(answer['formatted'])
            print("━" * 60 + "\n")


# 处理完成提示
print("✅ 所有内容处理完成！")