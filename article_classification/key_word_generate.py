import csv
import re
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.shared import RGBColor
from docx.oxml import OxmlElement, ns
from datetime import datetime
client = OpenAI(api_key="<API_KEY>", base_url="<BASE_URL>")
# 读取模板和示例（只需一次）
with open("IT_examples.txt", "r", encoding="utf-8") as example:
    example_content = example.read()

with open("IT_requirement.txt", "r", encoding="utf-8") as requirement:
    tech_requirement = requirement.read()

with open("key_word.txt", "r", encoding="utf-8") as key_word:
    key_words = key_word.read()
# 初始化文档设置
def init_document():
    doc = Document()
    section = doc.sections[0]

    # A4纸张设置
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # 页边距
    margins = section
    margins.left_margin = Cm(3.18)
    margins.right_margin = Cm(3.18)
    margins.top_margin = Cm(2.54)
    margins.bottom_margin = Cm(2.54)

    # 行网格设置
    sectPr = section._sectPr
    docGrid = OxmlElement('w:docGrid')
    docGrid.set(qn('w:type'), 'lines')
    docGrid.set(qn('w:linePitch'), '312')
    sectPr.append(docGrid)

    return doc


# 处理单篇文章
def add_article(doc, answer):
    """ 根据结构化answer添加文章 """
    # 解包数据
    title = answer['metadata']['title']
    content = answer['content']


    # ===== 标题部分 =====
    title_para = doc.add_paragraph()
    # Wingdings符号
    symbol_run = title_para.add_run(chr(184))
    symbol_run.font.name = "Wingdings 2"
    symbol_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Wingdings 2')
    symbol_run.font.color.rgb = RGBColor(0, 0, 0)  # 新增颜色设置
    # 标题文本
    title_run = title_para.add_run(f" {title}")
    title_run.font.name = "仿宋"
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色
    # 段落格式设置
    title_para.style = 'Heading 2'  # 通过样式名称设置大纲级别
    title_para.paragraph_format.line_spacing = 1.16  # 新增行距设置
    title_para.paragraph_format.space_after = Pt(8)

    # ===== 正文处理 =====
    paragraphs = [p.strip() for p in content.split('\n') if p.strip()][:2]
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.16  # 明确设置行距
        p.paragraph_format.first_line_indent = Pt(28)  # 首行缩进
        run = p.add_run(para)
        run.font.name = "仿宋"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(8)

    # 实时保存
    doc.save("./result/xinzhiyuan_keywordgenerate.docx")


# 主处理流程
doc = init_document()
seen_titles = set()
# 处理CSV文件
with open("../csv_folder/xinzhiyuan.csv", "r", encoding="utf-8") as csvfile:
    csv_reader = csv.reader(csvfile)

    while True:
        try:
            # 读取表头行+数据行
            headers = next(csv_reader)
            data = next(csv_reader)
        except StopIteration:
            break  # 文件结束

        # 构建数据字典（跳过空行检查）
        if len(headers) == 4 and len(data) == 4:
            row = dict(zip(headers, data))

            # ==== 重复标题检测 ====
            if row['title'] in seen_titles:
                print(f"跳过：{row['title']} （标题重复）")
                continue
            else:
                seen_titles.add(row['title'])

            # ==== 新增内容筛选 ====
            check_response = client.chat.completions.create(
                model="YOUR_MODEL",
                messages=[
                    {
                        "role": "system",
                        "content": f"""你是一个严谨的内容审核助手，专门判断一段科技新闻是否属于企业或者高校、研究所已经或者将要发布的科技成果，并且这些成果已经或即将产生经济效益。
                        
            附加要求：
                - 内容应聚焦信息技术（IT）相关领域，包括狭义 IT（如软件、算法、平台）和广义 IT（如人工智能、大模型、脑机接口、数字人等融合技术）。非技术类或与 IT 无明显关联的成果不在考虑范围内。
                - 请结合语义理解与关键词命中，进行综合判断。

            以下是一些关键词示例（不限于字面匹配，意思相近也可视为匹配）：
            {key_words}

            判断标准如下：
            1. 是否由企业、高校或科研机构主导或推动；
            2. 是否体现为一项具体科技成果（包括但不限于新技术、新产品、AI模型、demo展示、原型系统等）；
            3. 是否存在落地、转化、应用，或存在明确的经济潜力（即便未明说订单、盈利，也可以从产品演示、商业场景、技术壁垒等内容中合理推断）；
            4. 是否具备 IT 属性，且与关键词内容高度相关，可语义匹配；
            5. 是否至少覆盖五类关键词中任意三类（不要求五类全覆盖，但需体现多方面特征）。

            请根据上面的标准判断以下内容是否符合要求。只需回答“是”或“否”，不要添加解释。
            """
                    },
                    {
                        "role": "user",
                        "content": f"""请判断以下内容是否符合标准：
            内容如下：
            {row['passage']}"""
                    }
                ],
                temperature=0.0
            )

            judgement = check_response.choices[0].message.content.strip()
            if judgement != "是":
                print(f"跳过：{row['title']} （筛选未通过）")
                continue

            # 构建查询内容
            art_content = f"这是待处理的文章，请按模板处理：\n{row['passage']}\n。"

            # API调用
            response = client.chat.completions.create(
                model="YOUR_MODEL",
                messages=[
                    {"role": "system", "content": "你是一个专业的文本处理助理，能够完全理解并执行我提出的文本处理要求。"},
                    {"role": "user",
                     "content": f"我是一个科技类刊物的主编，我现在有一篇待处理的文章、几篇处理好的文章（主要是参考格式和语言风格）以及一个用于处理文章的要求模板，作为我的助理，你的任务是按照要求模板和示例来处理文章，并一定要保证处理结果的准确性。这是第三篇示例文章：\n{example_content}\n。"},
                    {"role": "assistant",
                     "content": "我已经学习了这几篇文章和对应的处理结果，我会学习每篇文章的处理方法，接下来请提供要求模板。"},
                    {"role": "user", "content": tech_requirement},  # 省略需求描述，可以从你的代码中复制需求部分
                    {"role": "assistant",
                     "content": "明白了，这是你给的处理要求模板，我将会严格遵循这个要求模板处理文章。接下来，请提供你需要我处理的文章，我会根据模板和示例的处理方式来进行处理，并严格保证处理结果的准确无误。"},
                    {"role": "user", "content": art_content}
                    # {"role": "user", "content": "在这之前，请复述一下三个示例的内容。"}
                ],
                temperature=1.1,
                top_p=0.9
            )
            # 构建结构化answer对象
            answer = {
                "metadata": {
                    "title": row['title']
                },
                "content": response.choices[0].message.content,
                "formatted": f"""\
            ▌标题：{row['title']}
            ────────────────────────────
            ▌生成内容：
            {response.choices[0].message.content}
            ────────────────────────────
            """
            }
            # 使用answer调用函数
            add_article(doc=doc,
                        answer=answer)  # 修改函数参数为接收answer
            # 使用格式化输出
            print(answer['formatted'])
            print("━" * 60 + "\n")


# 处理完成提示
print("✅ 所有内容处理完成！")